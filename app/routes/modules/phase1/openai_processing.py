
from docx.text.paragraph import Paragraph
from docx.table import Table
import json
import io
import os
import openai
from docx import Document
from openpyxl import load_workbook
# from google.colab import files
from openpyxl.styles import Alignment
from dotenv import load_dotenv
from openpyxl.styles import Font, PatternFill
import re
import tempfile


# Load environment variables from .env file
load_dotenv()

openai.api_key = os.getenv("OPENAI_API_KEY")

client = openai.OpenAI()

# Function: Read full document text


def extract_content_with_openai(docx_file):
    doc = Document(docx_file)
    full_text = "\n".join([para.text.strip()
                          for para in doc.paragraphs if para.text.strip()])

    print("🧠 Processing document with OpenAI...")

    prompt = f"""
You are given the full text of a Word document. Act as a senior bidding tender document analyst and your task is to extract the structure into a JSON array where each item contains:

- header (main section heading), Please keep the numbering etc shown in the document.
- subheader (if applicable), please keep the numbering etc shown in the document.
- requirements (list of extracted requirements under the section or subheader), it should not be rephrase or reworded, it should just be the same from the document to avoid confusions. Also keep the numbering etc same.
- page_limit (if mentioned in the text, otherwise 0)

Return only **valid JSON** in this format:

[
  {{
    "header": "Header Title",
    "subheader": "Subsection Title or null",
   "requirements": [
      "1) Requirement one",
      "2) Requirement two",
      "(a) Sub requirement",
      "(b) Another sub requirement"
    ],
    "page_limit": "2"
  }},
  ...
]

[
  {{
    "header": "Appendix A – Tender Submission Requirements",
    "subheader": "Annexure 1",
    "requirements": [
      "1) Requirement one",
      "2) Requirement two",
      "(a) Sub requirement",
      "(b) Another sub requirement"
    ],
    "page_limit": "0"
  }}
]

RULES:
- Preserve exact numbering and lettering (like "1)", "(a)", etc.)
- Keep the original punctuation and structure from the document
- Only respond with a **valid JSON array**
- Do NOT include any commentary, markdown, or natural language outside the JSON

Here is the document content:

{full_text}
"""

    try:
        response = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="gpt-4o-mini",
        )

        raw_content = response.choices[0].message.content.strip()
        # Truncate for display
        print("🔎 Raw GPT response:\n", raw_content[:1000])

        if raw_content.startswith("```json"):
            raw_content = raw_content[7:]
        if raw_content.endswith("```"):
            raw_content = raw_content[:-3]

        structured_data = json.loads(raw_content)
        return structured_data

    except Exception as e:
        print("❌ Error parsing GPT response:", str(e))
        return []


# Function: Excel formatting
def apply_wrap_text(cell):
    cell.alignment = Alignment(wrap_text=True)


def break_text_into_lines(text, max_characters=50):
    words = text.split(' ')
    lines = []
    current_line = []
    for word in words:
        if len(' '.join(current_line + [word])) > max_characters:
            lines.append(' '.join(current_line))
            current_line = [word]
        else:
            current_line.append(word)
    if current_line:
        lines.append(' '.join(current_line))
    return '\n'.join(lines)


def add_excel_with_sections(sections, excel_file):
    # Load the Excel file
    wb = load_workbook(excel_file)
    ws = wb.active

    # Headers
    ws["A1"] = "Header"
    ws["B1"] = "Subheader"
    ws["C1"] = "Requirements"
    ws["D1"] = "Page Limit"
    for cell in ["A1", "B1", "C1", "D1"]:
        apply_wrap_text(ws[cell])

    row_num = 2

    for item in sections:
        header = item.get("header", "")
        subheader = item.get("subheader", "")
        requirements = "\n".join(item.get("requirements", []))
        page_limit = item.get("page_limit", "0")

        # Column A - Header
        cell = ws[f"A{row_num}"]
        cell.value = break_text_into_lines(header)
        apply_wrap_text(cell)

        # Column B - Subheader
        cell = ws[f"B{row_num}"]
        cell.value = break_text_into_lines(subheader or "N/A")
        apply_wrap_text(cell)

        # Column C - Requirements
        cell = ws[f"C{row_num}"]
        cell.value = break_text_into_lines(requirements)
        apply_wrap_text(cell)

        # Column D - Page Limit
        cell = ws[f"D{row_num}"]
        cell.value = page_limit
        apply_wrap_text(cell)
        row_num += 1

    wb.save(excel_file)
    return excel_file


def extract_tables_from_docx_usingpydocx(word_file):
    tmp_path = None
    word_file.seek(0)
    print("are you here !")

    try:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(word_file.read())
            tmp_path = tmp.name
            print("you are inside the tempfile")

        print('you are inside the try:')
        doc = Document(tmp_path)
        tables_data = []
        print("number of tables in document is", len(doc.tables))

        for index, table in enumerate(doc.tables, start=1):
            table_content = []
            for row in table.rows:
                wrapped_row = []
                for cell in row.cells:
                    text = cell.text.strip()
                    wrapped_row.append({
                        "text": text,
                        "column_header": bool(text),  # naïve header logic
                        "row_header": bool(text),     # naïve header logic
                    })
                table_content.append(wrapped_row)

            tables_data.append({
                "heading": f"Table {index}",
                "table": table_content
            })

        print("table data is:")
        print(tables_data)
        return tables_data

    finally:
        # Clean up: Delete the temporary file
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)


def extract_tables_with_headings_and_context(word_file):
    tmp_path = None
    word_file.seek(0)
    print("are you here !")

    try:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(word_file.read())
            tmp_path = tmp.name
            print("you are inside the tempfile")

        print('you are inside the try:')
        doc = Document(tmp_path)
        tables_data = []

        unknown_idx = 1
        # Helper function to iterate through paragraphs and tables in order

        def iter_block_items(parent):
            for child in parent.element.body.iterchildren():
                if child.tag.endswith('}p'):
                    yield Paragraph(child, parent)
                elif child.tag.endswith('}tbl'):
                    yield Table(child, parent)

        all_elements = list(iter_block_items(doc))
        current_heading = None

        for idx, element in enumerate(all_elements):
            if isinstance(element, Paragraph):
                style = element.style.name if element.style else ''
                if style.startswith('Heading'):
                    current_heading = element.text.strip()

            elif isinstance(element, Table):
                table_content = []
                for row in element.rows:
                    row_data = []
                    for cell in row.cells:
                        text = cell.text.strip()
                        row_data.append({
                            "text": text,
                            "column_header": bool(text),
                            "row_header": bool(text)
                        })
                    table_content.append(row_data)

                # Capture up to 2 paragraphs before and after the table for context
                before_context = []

                # Look back for up to 2 paragraphs
                i = idx - 1
                while i >= 0 and len(before_context) < 4:
                    prev_elem = all_elements[i]
                    if isinstance(prev_elem, Paragraph):
                        before_context.insert(0, prev_elem.text.strip())
                    i -= 1

                pred_heading = all_elements[idx-1].text.strip()
                if (len(pred_heading) > 45 or len(pred_heading) == 0):
                    curr_heading = f"Unknown {unknown_idx}"
                    unknown_idx += 1
                else:
                    curr_heading = pred_heading

                tables_data.append({
                    "heading": curr_heading,
                    "table": table_content,
                    "before_context": before_context})

        return tables_data
    finally:
        # Clean up: Delete the temporary file
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)

# def extract_tables_from_docx(word_file):
#     tmp_path = None
#     word_file.seek(0)

#     try:
#         # Create a temporary file
#         with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
#             tmp.write(word_file.read())
#             tmp_path = tmp.name

#         converter = Converter()
#         json_format = converter.convert_docx_to_json(tmp_path)
#         _tables = json_format["tables"]

#         structured_tables = []
#         for idx, tbl in enumerate(docling_tables, start=1):
#             # 1) choose a heading
#             caps = tbl.get("captions", [])
#             if caps:
#                 heading = caps[0] if isinstance(
#                     caps[0], str) else caps[0].get("text", "")
#             else:
#                 heading = f"Table {idx}"

#             # 2) wrap each cell dict with just the bits we need
#             grid = tbl["data"]["grid"]
#             table = []
#             for row in grid:
#                 wrapped_row = []
#                 for cell in row:
#                     wrapped_row.append({
#                         "text":           cell.get("text", "").strip(),
#                         "column_header":  bool(cell.get("column_header", False)),
#                         "row_header":     bool(cell.get("row_header",    False)),
#                     })
#                 table.append(wrapped_row)

#             structured_tables.append({
#                 "heading": heading,
#                 "table":   table
#             })

#         return structured_tables
#     finally:
#         # Clean up: Delete the temporary file
#         if tmp_path and os.path.exists(tmp_path):
#             os.remove(tmp_path)


def sanitize_sheet_title(title: str) -> str:
    # Excel sheet titles: max 31 chars, cannot contain : \ / ? * [ ]
    safe = re.sub(r'[:\\\/\?\*\[\]]', '_', title)
    return safe[:31]


def add_excel_with_tables(tables, excel_file):

    # Load the Excel file
    wb = load_workbook(excel_file)
    ws = wb.active

    header_font = Font(bold=True)

    for tbl in tables:
        title = sanitize_sheet_title(tbl["heading"])
        ws = wb.create_sheet(title=title)

        for r, row in enumerate(tbl["table"], start=1):
            for c, cell in enumerate(row,   start=1):
                # should print something like {'text': '...', ...}
                text = cell["text"]
                # write the TEXT only, not the whole dict
                excel_cell = ws.cell(row=r, column=c, value=text)
                if cell["column_header"] or cell["row_header"]:
                    excel_cell.font = header_font
                    excel_cell.fill = PatternFill("lightUp", fgColor="FFC400")
                # Extend the width of the excel cell to fit the text fully
                from openpyxl.utils import get_column_letter
                ws.column_dimensions[get_column_letter(c)].width = max(
                    len(str(text)) + 2, ws.column_dimensions[get_column_letter(c)].width)

    # # Create an in-memory buffer to store the modified Excel file
    # excel_output = io.BytesIO()
    wb.save(excel_file)
    # # Reset the pointer to the start of the buffer
    # excel_output.seek(0)
    return excel_file

    # wb.save(output_path)
    # print(f"✅ Excel file saved to: {output_path}")
